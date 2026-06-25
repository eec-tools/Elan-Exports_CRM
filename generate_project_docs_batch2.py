#!/usr/bin/env python3
"""
Generate Word documents for projects 5-8:
Sales, Product-Management, Marketing, Legal
Uses keyword-based conversation matching since all 4 share the same creation date.
"""

import json
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/harshpatel/Desktop/Elan-Exports_CRM/data-6d74feb6-cd86-4d3b-a609-ee0f3f215078-1781522434-c91936d3-batch-0000"
OUT_DIR = "/Users/harshpatel/Desktop/Elan-Exports_CRM"

# ── keyword matching config ────────────────────────────────────────────────────

PROJECT_KEYWORDS = {
    "Sales": [
        'sales', 'buyer', 'crm', 'lead', 'prospect', 'sourcing', 'vendor',
        'pricing', 'deal', 'client', 'supplier', 'export', 'import', 'enquiry',
        'rfq', 'quotation', 'purchase', 'order', 'negotiate', 'outreach',
        'pipeline', 'revenue', 'target', 'commission', 'distributor', 'invoice',
        'manufacturer', 'bulk', 'commodit', 'textile', 'lentil', 'rice', 'corn',
        'fortified', 'pulse', 'cotton', 'african', 'egypt', 'excel', 'csv',
        'vendor list', 'checking vendor', 'eec sourcing', 'eec - sourcing',
        'product enquiry', 'tiered sourcing', 'buyer qualification', 'buyer-facing',
        'ai agent.*sales', 'setting up eec sales', 'gtm', 'go-to-market'
    ],
    "Product-Management": [
        'product', 'pm ', 'roadmap', 'prd', 'feature', 'saas', 'spec',
        'mvp', 'sprint', 'backlog', 'user stor', 'requirement', 'milestone',
        'crm.*build', 'build.*crm', 'react', 'dashboard', 'application',
        'agent.*setup', 'agent.*structure', 'crm report', 'ceo dashboard',
        'crm.*redesign', 'eec.*plan', 'new plan', 'aws', 'deployment',
        'security', 'migration', 'infrastructure', 'excel.*export', 'tool link',
        'deal value', 'timing constraint', 'bot setup', 'telegram',
        'api', 'build', 'develop', 'engineer', 'architect', 'system',
        'comprehensive react', 'notion integrat', 'ai agent.*eec',
        'orchestration', 'ai.*agent.*build', 'llm', 'council'
    ],
    "Marketing": [
        'marketing', 'instagram', 'brand', 'content', 'social', 'reel',
        'campaign', 'audience', 'engagement', 'seo', 'advertisement', 'post',
        'growth', 'faceless', 'video', 'script', 'visual', 'creative',
        'promot', 'influencer', 'caption', 'hashtag', 'strategy.*market',
        'market.*strategy', 'brand.*name', 'naming', 'export consultancy.*brand',
        'brand.*export', 'notebooklm', 'presentation', 'mind map',
        'founding brief', 'identity framework', 'blog post', 'textile.*import',
        'source research', 'market research', 'claude.*marketing'
    ],
    "Legal": [
        'legal', 'contract', 'compliance', 'nda', 'agreement', 'ip',
        'intellectual property', 'law ', 'clause', 'terms', 'liabilit',
        'dispute', 'protect', 'rights', 'patent', 'trademark', 'copyright',
        'transfer agreement', 'code.*nda', 'protecting.*code', 'securing',
        'confidentiaI', 'confidential', 'auspicious', 'email address.*select',
        'university tuition', 'compound interest', 'cbdc', 'digital currency',
        'money', 'financial', 'air miles', 'credit card', 'finding.*council',
        'claude council', 'scheduling tasks', 'disconnected.*number',
        'singapore number', 'meal plan', 'home', 'life hq', 'my life',
        'linkedin.*resume', 'resume.*evaluation', 'linkedin.*profile', 'job description',
        'intern', 'delegating', 'notion.*space', 'ss -'
    ],
}

# Conversations already used in batch 1 (date ranges already covered)
ALREADY_USED_DATES = set()  # We'll include ALL — better to have some overlap than miss


def score_conversation(conv, keywords):
    """Score a conversation against keyword list. Higher = better match."""
    name = (conv.get('name', '') or conv.get('summary', '') or '').lower()
    score = 0
    for kw in keywords:
        if kw.lower() in name:
            score += 3  # Name match is strongest signal

    # Check first human message text
    msgs = conv.get('chat_messages', [])
    for msg in msgs[:4]:
        if msg.get('sender') == 'human':
            text = (msg.get('text', '') or '').lower()
            for kw in keywords:
                if kw.lower() in text:
                    score += 1
    return score


def assign_conversations_to_projects(conversations, project_names):
    """Assign each conversation to exactly one project based on best keyword match."""
    assignments = {name: [] for name in project_names}
    unassigned = []

    for conv in conversations:
        scores = {
            name: score_conversation(conv, PROJECT_KEYWORDS[name])
            for name in project_names
        }
        best = max(scores, key=scores.get)
        if scores[best] > 0:
            assignments[best].append((conv, scores[best]))
        else:
            unassigned.append(conv)

    # Put unassigned conversations in Sales (catch-all)
    for conv in unassigned:
        assignments["Sales"].append((conv, 0))

    # Sort each project's conversations by date
    for name in project_names:
        assignments[name] = [
            c for c, _ in sorted(assignments[name], key=lambda x: x[0].get('created_at', ''))
        ]

    return assignments


# ── formatting helpers (same as batch 1) ─────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_colored_heading(doc, text, level, color_hex="1F497D"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_bordered_left(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'right', 'bottom', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:color'), color)
    tcBorders.append(left)
    tcPr.append(tcBorders)


def add_message_block(doc, sender, text, timestamp):
    is_human = sender == "human"

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]

    if is_human:
        set_cell_bg(cell, "E8F4FD")
        add_bordered_left(cell, "2E74B5")
        label_color = "2E74B5"
        label = "YOU"
    else:
        set_cell_bg(cell, "F0F9F0")
        add_bordered_left(cell, "27AE60")
        label_color = "1E8449"
        label = "CLAUDE"

    p_header = cell.paragraphs[0]
    label_run = p_header.add_run(f"{label}   {timestamp}")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor.from_string(label_color)

    lines = text.strip().split('\n')
    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_cover_page(doc, project, conv_count, msg_count):
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(project['name'])
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor.from_string("1F497D")

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Claude AI Project Export")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor.from_string("666666")
    sub_run.italic = True

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    info_items = [
        ("Project Name",          project['name']),
        ("Created",               project.get('created_at', '')[:10]),
        ("Last Updated",          project.get('updated_at', '')[:10]),
        ("Knowledge Base Files",  str(len(project.get('docs', [])))),
        ("Conversations Included", str(conv_count)),
        ("Total Messages",        str(msg_count)),
        ("Exported By",           "Elan (sales@elanexports.com)"),
        ("Export Date",           "2026-06-18"),
    ]

    table = doc.add_table(rows=len(info_items), cols=2)
    table.style = 'Table Grid'
    for i, (label, val) in enumerate(info_items):
        row = table.rows[i]
        set_cell_bg(row.cells[0], "D6E4F0")
        set_cell_bg(row.cells[1], "EBF3FB")
        r = row.cells[0].paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("1F497D")
        row.cells[1].paragraphs[0].add_run(val).font.size = Pt(10)

    doc.add_page_break()


def build_project_info(doc, project):
    add_colored_heading(doc, "Project Overview", 1, "1F497D")
    add_horizontal_rule(doc)

    description = project.get('description', '') or 'No description provided.'
    desc_p = doc.add_paragraph(description)
    desc_p.paragraph_format.space_before = Pt(6)
    desc_p.paragraph_format.space_after = Pt(12)
    for run in desc_p.runs:
        run.font.size = Pt(11)

    prompt = project.get('prompt_template', '').strip()
    if prompt and prompt not in ("''", '""', ""):
        add_colored_heading(doc, "System Prompt (Custom Instructions)", 2, "2E74B5")
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        set_cell_bg(cell, "FFF8DC")
        add_bordered_left(cell, "E67E22")
        p = cell.paragraphs[0]
        r = p.add_run(prompt)
        r.font.size = Pt(10)
        r.font.name = "Courier New"
        doc.add_paragraph()


def build_docs_section(doc, project):
    docs = project.get('docs', [])
    if not docs:
        add_colored_heading(doc, "Knowledge Base", 1, "1F497D")
        add_horizontal_rule(doc)
        p = doc.add_paragraph("No knowledge base files were uploaded to this project.")
        p.runs[0].font.color.rgb = RGBColor.from_string("888888")
        p.runs[0].italic = True
        doc.add_paragraph()
        return

    add_colored_heading(doc, f"Knowledge Base ({len(docs)} file{'s' if len(docs)!=1 else ''})", 1, "1F497D")
    add_horizontal_rule(doc)

    for i, doc_item in enumerate(docs):
        filename = doc_item.get('filename', f'Document {i+1}')
        content = doc_item.get('content', '')
        created = doc_item.get('created_at', '')[:10] if doc_item.get('created_at') else ''

        file_heading = doc.add_paragraph()
        file_heading.paragraph_format.space_before = Pt(12)
        icon_run = file_heading.add_run(f"FILE:  {filename}")
        icon_run.bold = True
        icon_run.font.size = Pt(13)
        icon_run.font.color.rgb = RGBColor.from_string("2E74B5")

        if created:
            meta_p = doc.add_paragraph(f"Uploaded: {created}  |  Size: {len(content):,} characters")
            meta_p.runs[0].font.size = Pt(9)
            meta_p.runs[0].font.color.rgb = RGBColor.from_string("888888")

        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        set_cell_bg(cell, "F8F9FA")

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'right', 'bottom']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), 'CCCCCC')
            tcBorders.append(border)
        tcPr.append(tcBorders)

        lines = content.strip().split('\n')
        first = True
        for line in lines:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        doc.add_paragraph()

    doc.add_page_break()


def build_conversations_section(doc, conversations, matching_note):
    total_msgs = sum(len(c.get('chat_messages', [])) for c in conversations)

    add_colored_heading(
        doc,
        f"Conversations ({len(conversations)} conversations, {total_msgs} messages)",
        1, "1F497D"
    )

    note_p = doc.add_paragraph()
    note_run = note_p.add_run(f"Note: {matching_note}")
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor.from_string("888888")
    note_p.paragraph_format.space_after = Pt(8)

    add_horizontal_rule(doc)

    if not conversations:
        p = doc.add_paragraph("No conversations matched this project's topic keywords.")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor.from_string("888888")
        return

    for conv_idx, conv in enumerate(conversations):
        name = conv.get('name', '') or conv.get('summary', '') or f"Conversation {conv_idx + 1}"
        created = conv.get('created_at', '')[:10]
        msgs = conv.get('chat_messages', [])

        conv_header = doc.add_paragraph()
        conv_header.paragraph_format.space_before = Pt(16)
        h_run = conv_header.add_run(f"---  {name}")
        h_run.bold = True
        h_run.font.size = Pt(12)
        h_run.font.color.rgb = RGBColor.from_string("34495E")

        meta_p = doc.add_paragraph(f"Date: {created}  |  Messages: {len(msgs)}")
        meta_p.runs[0].font.size = Pt(9)
        meta_p.runs[0].font.color.rgb = RGBColor.from_string("999999")

        real_msgs = [m for m in msgs if m.get('text', '').strip() or m.get('content')]

        if not real_msgs:
            p = doc.add_paragraph("(No message content available in this export)")
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor.from_string("AAAAAA")
            p.runs[0].font.size = Pt(9)
            continue

        for msg in real_msgs:
            sender = msg.get('sender', 'unknown')
            text = msg.get('text', '').strip()
            ts = msg.get('created_at', '')[:16].replace('T', ' ')

            if not text:
                content_blocks = msg.get('content', [])
                texts = []
                for block in content_blocks:
                    if isinstance(block, dict):
                        if block.get('type') == 'text':
                            texts.append(block.get('text', ''))
                        elif 'text' in block:
                            texts.append(block['text'])
                text = '\n'.join(texts).strip()

            if not text:
                continue

            if len(text) > 8000:
                text = text[:8000] + f"\n\n[... {len(text) - 8000:,} more characters truncated ...]"

            add_message_block(doc, sender, text, ts)

        add_horizontal_rule(doc)


def generate_doc(project, conversations, matching_note, output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    conv_count = len(conversations)
    msg_count = sum(len(c.get('chat_messages', [])) for c in conversations)

    build_cover_page(doc, project, conv_count, msg_count)
    build_project_info(doc, project)
    build_docs_section(doc, project)
    build_conversations_section(doc, conversations, matching_note)

    doc.save(output_path)
    print(f"  Saved: {output_path}")


# ── main ───────────────────────────────────────────────────────────────────────

def main():
    project_files = sorted(os.listdir(f"{BASE}/projects"))
    all_projects = []
    for pf in project_files:
        with open(f"{BASE}/projects/{pf}") as f:
            all_projects.append(json.load(f))

    # Projects 5-8 (index 4-7): Sales, Product-Management, Marketing, Legal
    target_projects = all_projects[4:8]
    project_names = [p['name'] for p in target_projects]
    print(f"Target projects: {project_names}")

    with open(f"{BASE}/conversations.json") as f:
        all_conversations = json.load(f)

    print(f"Total conversations in export: {len(all_conversations)}")

    # Assign conversations to projects by keyword scoring
    assignments = assign_conversations_to_projects(all_conversations, project_names)

    matching_note = (
        "These conversations were matched to this project using keyword/topic scoring "
        "because all 4 projects (Sales, Product-Management, Marketing, Legal) were created "
        "on the same date (Apr 29, 2026) and the Claude export does not explicitly link "
        "conversations to projects. Each conversation was assigned to the project whose "
        "keywords best matched the conversation title and opening message."
    )

    for project in target_projects:
        name = project['name']
        conversations = assignments[name]
        safe_name = name.replace('/', '-').replace('\\', '-').replace(':', '-').replace('–', '-').strip()
        output_path = f"{OUT_DIR}/{safe_name}.docx"

        msg_count = sum(len(c.get('chat_messages', [])) for c in conversations)
        print(f"\nGenerating: {safe_name}")
        print(f"  Conversations: {len(conversations)}, Messages: {msg_count}")

        generate_doc(project, conversations, matching_note, output_path)

    print("\nAll 4 documents generated.")


if __name__ == "__main__":
    main()
