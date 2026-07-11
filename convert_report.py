#!/usr/bin/env python3
"""
BUYER_AGENT_REPORT.docx — Elan Exports Buyer Discovery Agent Report
Sections: Cover + 1 (Goal) + 2 (Tech Stack) + 3 (Methods) + Task Tracker
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX_FILE = "BUYER_AGENT_REPORT.docx"

# ── Colours ───────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x0F, 0x2B, 0x5B)
MED_BLUE   = RGBColor(0x1A, 0x4A, 0x9E)
LIGHT_BLUE = RGBColor(0xEF, 0xF6, 0xFF)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_BG    = RGBColor(0xF8, 0xFA, 0xFF)
RED_BG     = RGBColor(0xFF, 0xF1, 0xF2)
RED_TEXT   = RGBColor(0x7F, 0x1D, 0x1D)
AMBER_BG   = RGBColor(0xFF, 0xFB, 0xEB)
AMBER_TEXT = RGBColor(0x78, 0x35, 0x0F)
CODE_TEXT  = RGBColor(0xE2, 0xE8, 0xF0)
BODY_TEXT  = RGBColor(0x37, 0x41, 0x51)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

doc.styles["Normal"].font.name  = "Calibri"
doc.styles["Normal"].font.size  = Pt(11)
doc.styles["Normal"].font.color.rgb = BODY_TEXT

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_inline_formatted(para, text):
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = MED_BLUE
        else:
            if part:
                para.add_run(part)

section_num = [0]

def add_section_heading(title_text):
    section_num[0] += 1
    para = doc.add_heading(f"{section_num[0]}.  {title_text}", level=1)
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)
    for run in para.runs:
        run.font.name  = "Calibri"
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = DARK_BLUE
    return para

def add_h3(text):
    para = doc.add_heading(text, level=2)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)
    for run in para.runs:
        run.font.name  = "Calibri"
        run.font.size  = Pt(12)
        run.font.bold  = True
        run.font.color.rgb = MED_BLUE
    return para

def add_body(text, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(space_after)
    add_inline_formatted(para, text)
    for run in para.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return para

def add_bullet(text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    add_inline_formatted(para, text)
    for run in para.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return para

def add_code_block(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(8)
    run = para.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = CODE_TEXT
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1E293B")
    pPr.append(shd)
    para.paragraph_format.left_indent  = Inches(0.2)
    para.paragraph_format.right_indent = Inches(0.2)
    return para

def add_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(9.5)
        run.font.color.rgb = WHITE
        run.font.name  = "Calibri"
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri + 1]
        bg = GREY_BG if ri % 2 == 1 else WHITE
        for ci, cell_text in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.clear()
            add_inline_formatted(p, str(cell_text))
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def add_callout(text, style="blue"):
    colors = {
        "blue":  (LIGHT_BLUE, MED_BLUE),
        "red":   (RED_BG,    RED_TEXT),
        "amber": (AMBER_BG,  AMBER_TEXT),
    }
    bg, fg = colors.get(style, colors["blue"])
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(8)
    para.paragraph_format.left_indent  = Inches(0.15)
    add_inline_formatted(para, text)
    for run in para.runs:
        run.font.name  = "Calibri"
        run.font.size  = Pt(10.5)
        run.font.color.rgb = fg
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{bg[0]:02X}{bg[1]:02X}{bg[2]:02X}")
    pPr.append(shd)
    return para

def spacer(pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.LEFT
tag_run = tag.add_run("INTERNAL TECHNICAL REPORT")
tag_run.font.size  = Pt(9)
tag_run.font.bold  = True
tag_run.font.color.rgb = MED_BLUE
tag_run.font.name  = "Calibri"
tag.paragraph_format.space_before = Pt(0)
tag.paragraph_format.space_after  = Pt(4)

title = doc.add_heading("Buyer Discovery Agent", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.runs[0]
run.font.size  = Pt(28)
run.font.bold  = True
run.font.color.rgb = DARK_BLUE
run.font.name  = "Calibri"
title.paragraph_format.space_after = Pt(2)

subtitle = doc.add_paragraph("Development Report")
r2 = subtitle.runs[0] if subtitle.runs else subtitle.add_run("Development Report")
r2.font.size  = Pt(16)
r2.font.color.rgb = MED_BLUE
r2.font.name  = "Calibri"
subtitle.paragraph_format.space_after = Pt(2)

subline = doc.add_paragraph('Agent 1 — "Discover & Rank"  ·  Elan Exports & Co. CRM')
r3 = subline.runs[0] if subline.runs else subline.add_run('Agent 1 — "Discover & Rank"  ·  Elan Exports & Co. CRM')
r3.font.size  = Pt(12)
r3.font.color.rgb = RGBColor(0x60, 0x7D, 0xA8)
r3.font.name  = "Calibri"
subline.paragraph_format.space_after = Pt(14)

meta_tbl = doc.add_table(rows=1, cols=4)
meta_tbl.style = "Table Grid"
meta_items = [
    ("Company",  "Elan Exports & Co. (EEC)"),
    ("Date",     "July 2026"),
    ("Agent",    "Agent 1 — Discover & Rank"),
    ("Status",   "⚠ Partial — Email Tools Need Paid Plan"),
]
for i, (label, value) in enumerate(meta_items):
    cell = meta_tbl.cell(0, i)
    set_cell_bg(cell, GREY_BG)
    p = cell.paragraphs[0]
    p.clear()
    lbl = p.add_run(label + "\n")
    lbl.font.size  = Pt(7.5)
    lbl.font.bold  = True
    lbl.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    lbl.font.name  = "Calibri"
    val = p.add_run(value)
    val.font.size  = Pt(10)
    val.font.bold  = True
    val.font.color.rgb = DARK_BLUE
    val.font.name  = "Calibri"
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

spacer(10)

banner = doc.add_paragraph()
banner.paragraph_format.space_before = Pt(4)
banner.paragraph_format.space_after  = Pt(12)
b1 = banner.add_run("⚠  CURRENT STATUS:  ")
b1.bold = True
b1.font.color.rgb = AMBER_TEXT
b1.font.name = "Calibri"
b2 = banner.add_run(
    "The pipeline is fully built and runs end-to-end. Discovery (finding companies) works. "
    "The blocker is email finding — Snov.io prospect search and Apollo people search are both "
    "paywalled on free plans, and the companies we find (small EU/UK textile SMEs) are not "
    "indexed in Snov.io or Hunter's free database."
)
b2.font.color.rgb = AMBER_TEXT
b2.font.name = "Calibri"
b2.font.size = Pt(10.5)

pPr = banner._p.get_or_add_pPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:val"),   "clear")
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"),  f"{AMBER_BG[0]:02X}{AMBER_BG[1]:02X}{AMBER_BG[2]:02X}")
pPr.append(shd)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — GOAL
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading("What We Are Trying to Do")

add_body(
    "Elan Exports is an India-based sourcing intermediary in commodities and textiles. "
    "Agent 1 automatically finds companies in target countries (UK, Germany, France, UAE, etc.) "
    "that **import** the products Elan Exports can supply — and then finds the **direct email of "
    "the procurement / buying person** at that company so the sales team can reach out."
)

add_h3("What 'Real and Genuine' Contact Means")
add_table(
    ["We WANT", "We DO NOT WANT"],
    [
        ["`james.hartley@fabricimports.co.uk` — named buyer with procurement title",
         "`info@company.com` — general inbox, nobody specific reads it"],
        ["`procurement@whaleys-bradford.co.uk` — buying department email",
         "`sales@company.com` — goes to sales team, not procurement"],
        ["`sourcing@textilegroup.de` — sourcing department email",
         "`contact@company.com` — black-hole inbox"],
        ["Verified email that reaches the decision maker",
         "Generic mailboxes that reach nobody relevant"],
    ]
)

add_h3("Target Countries & Product Categories")
add_table(
    ["Countries", "Product Categories"],
    [["Germany, France, Netherlands, Italy, Spain, UK, UAE, Saudi Arabia, Qatar, Kuwait, Japan, Singapore, Australia",
      "Textiles, Organic Food, Seafood, Rice & Grains, Spices & Herbs, Pulses & Lentils"]]
)

add_h3("Target Company Profile")
add_bullet("**Textiles:** Fabric wholesalers, yarn importers, lining/interlining importers, denim distributors that buy raw material from Asian mills")
add_bullet("**Organic Food:** Bio food importers, organic produce wholesalers, health food distributors")
add_bullet("**NOT:** Fashion brands (Balmain, Zara, H&M), supermarket chains (Tesco, Carrefour), logistics/freight companies")
add_bullet("**Size:** 30–500 employees — small enough to have a reachable procurement email, large enough to be a real B2B buyer")


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — TECH STACK
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_section_heading("Technical Stack Built")

add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Backend",          "Node.js + TypeScript",            "Agent pipeline, API server"],
        ["Database",         "PostgreSQL + Prisma ORM",         "Store runs, companies, contacts"],
        ["LLM",              "Groq API (LLaMA-3.3-70B)",        "Company discovery + AI scoring"],
        ["Web Search",       "Firecrawl API",                   "Google web search + website scraping"],
        ["Lead Search",      "Apollo.io API",                   "Find buying managers by job title"],
        ["Email Discovery",  "Snov.io API",                     "Prospect search + domain email lookup"],
        ["Email Fallback",   "Hunter.io API",                   "Domain email search (secondary)"],
        ["Frontend",         "React + TypeScript + TailwindCSS", "Run dashboard, results table, drawer"],
        ["Real-time",        "React Query (5s polling)",         "Live progress during runs"],
    ]
)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — ALL METHODS TRIED
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_section_heading("All Methods Tried — Chronological History")

# --- Method 1 ---
add_h3("Method 1 — Apollo.io + Firecrawl + Hunter.io + Groq  [Result: 0 leads]")
add_body("**What we tried:** Apollo.io was the primary discovery engine. The plan: search for 'Procurement Manager' / 'Buying Director' at textile companies in the target country, enrich with Firecrawl, verify emails with Hunter, score with Groq.")
add_body("**What happened:** Apollo's people search API is completely blocked on the free plan:")
add_code_block(
    "[Apollo] Search error: {\n"
    "  error: 'api/v1/mixed_people/search is not accessible with this api_key on a free plan.',\n"
    "  error_code: 'API_INACCESSIBLE'\n"
    "}\n"
    "[Apollo] 0 unique end-buyer leads for Textiles in Germany"
)
add_callout("**Root cause:** Apollo free plan blocks all `/mixed_people/search` endpoints. The API key works for authentication but search is paywalled. Needs Basic plan (~$49/month).", "red")
spacer()

# --- Method 2 ---
add_h3("Method 2 — Kompass + Europages Directory Scraping  [Result: Wrong data type]")
add_body("**What we tried:** B2B trade directories (Kompass.com, Europages) list thousands of importing companies. The plan was to use Firecrawl to search `site:kompass.com textile importer UK` to get actual importer company pages.")
add_body("**What happened:** Google returned Kompass's own category listing pages — not individual company pages: `lu.kompass.com`, `gr.kompass.com`, `in.kompass.com`. These are regional directory pages, not company websites.")
add_callout("**Root cause:** The SKIP_DOMAINS filter used exact matching — `lu.kompass.com` was not caught by the `kompass.com` rule. Fixed by adding parent domain matching so all Kompass subdomains are auto-blocked.", "amber")
spacer()

# --- Method 3 ---
add_h3("Method 3 — Groq LLM as Primary Company Discovery  [Result: 0 scored — hallucinated domains]")
add_body("**What we tried:** Since Apollo was blocked, Groq's LLaMA-3.3-70B model was used as the primary discovery engine — asking it to name textile importers in target countries.")
add_body("**What happened:** The LLM returned plausible-sounding company names with completely invented domains:")
add_table(
    ["LLM Generated", "Reality"],
    [
        ["`alkhaleejtextiles.com`",   "Does not exist — DNS: NXDOMAIN"],
        ["`gulfyarntrading.com`",     "Does not exist — DNS: NXDOMAIN"],
        ["`emiratestextilemills.com`","Does not exist — DNS: NXDOMAIN"],
        ["`pioneertextilesfze.com`",  "Does not exist — DNS: NXDOMAIN"],
        ["`nationaltextilellc.com`",  "Does not exist — DNS: NXDOMAIN"],
    ]
)
add_callout("**Root cause: LLM Hallucination.** Language models generate names that sound real but don't exist. They are text pattern generators, not business directories. This is a fundamental limitation — no amount of prompt engineering fully fixes it.", "red")
add_body("**Also fixed:** For France Textiles, the LLM returned Balmain, SMCP, Agnes B — luxury fashion brands, not fabric importers. A `CATEGORY_CONFIG` system was added with `whatWeWant`, `whatWeAvoid`, `goodExamples`, and `badExamples` fields per category.")
spacer()

# --- Method 4 ---
add_h3("Method 4 — Snov.io Full Integration (Prospect Search + Domain Search)  [Result: 0 — paywalled]")
add_body("**What we tried:** Snov.io's prospect search endpoint directly searches their LinkedIn-indexed database for people by job title + country + industry. A complete Snov.io integration was built:")
add_bullet("`snovProspectSearch()` — search for 'Procurement Manager' + 'Wholesale/Textiles' + 'UK'")
add_bullet("`snovDomainSearch()` — find emails for a known company domain")
add_bullet("`snovVerifyEmail()` — async two-step email verification (submit → poll)")
add_bullet("Credit exhaustion guard — stops all Snov.io calls when 50 free credits run out")
add_bullet("OAuth2 token caching with 60-second early-refresh")
add_body("**What happened — Prospect Search:**")
add_code_block("[Snov] Prospect search unavailable on current plan (HTTP 404)")
add_callout("**Prospect search is completely paywalled.** The endpoint `POST /v1/get-prospects-with-target` does not exist for free users. Requires Snov.io Starter plan (~$30/month). The code is fully written and ready — just needs the plan upgrade.", "red")
add_body("**Credit exhaustion problem:** 10 prefixes × 0.5 credits × 25 companies = 125 credits needed, only 50 free. Fixed by adding a `_creditsExhausted` flag that stops all Snov.io calls once credits run out.")
spacer()

# --- Method 5 ---
add_h3("Method 5 — Named Person Email Gate  [Result: Higher quality, still 0]")
add_body("**Trigger:** The sales team said they don't want `info@` or `sales@` emails — they want a real named person in procurement who they can source through.")
add_body("**What changed:**")
add_bullet("**Removed:** Any scraping that returned generic mailboxes (info@, contact@)")
add_bullet("**Hard filter (Snov.io):** Only accepted if `firstName` field is non-empty — proves a named person, not a dept mailbox")
add_bullet("**Hard filter (Hunter):** Only accepted if `first_name` is non-empty AND title matches procurement keywords")
add_code_block(
    "// Only named persons accepted — no generic mailboxes\n"
    ".filter((e) => e.firstName && e.firstName.trim().length > 0)\n"
    ".filter((e) => isProcurementTitle(e.position))"
)
add_body("**Result:** Email quality improved when contacts are found. Still 0 results because Snov.io and Hunter have no data for the small companies being discovered.")
spacer()

# --- Method 6 ---
add_h3("Method 6 — Firecrawl First + DNS Validation  [Result: Real companies found — email still 0]")
add_body("**Problem identified:** The LLM was running first every time. Firecrawl was only triggered if the LLM returned 0 — which it never did because it always generated something (even if fake).")
add_body("**Fix — Swapped order:**")
add_bullet("**Before:** LLM → Firecrawl (only if LLM returns 0)")
add_bullet("**After:** Firecrawl → LLM (only if Firecrawl returns < 10 results)")
add_body("**DNS validation added:** Before calling Snov.io or Hunter, check if the domain resolves. Fake LLM domains are caught instantly without spending credits:")
add_code_block(
    "import { lookup as dnsLookup } from 'dns/promises';\n\n"
    "async function domainResolves(domain: string): Promise<boolean> {\n"
    "  try { await dnsLookup(domain); return true; }\n"
    "  catch { return false; }  // NXDOMAIN = domain doesn't exist\n"
    "}"
)
add_body("**UK Textiles run result:** Firecrawl found 39 companies — but many were wrong type: `ibisworld.com` (market research), `scribd.com` (document sharing), `lusha.com` (lead gen tool), `ukft.org` (trade association). Real companies found: `whaleys-bradford.ltd.uk`, `oddies-textiles.co.uk` — but not indexed in any email database.")
spacer()

# --- Method 7 ---
doc.add_page_break()
add_h3("Method 7 — SKIP_DOMAINS Expansion + Article Filter + Tier 3 Dept Email  [Current — Best Version]")
add_body("**Three improvements applied in the current version:**")

add_body("**Fix A — SKIP_DOMAINS expanded to 40+ entries with subdomain matching:**")
add_code_block(
    "function shouldSkipDomain(domain: string): boolean {\n"
    "  if (SKIP_DOMAINS.has(domain)) return true;\n"
    "  // lu.kompass.com → caught by kompass.com\n"
    "  const parts = domain.split('.');\n"
    "  for (let i = 1; i < parts.length - 1; i++) {\n"
    "    if (SKIP_DOMAINS.has(parts.slice(i).join('.'))) return true;\n"
    "  }\n"
    "  return false;\n"
    "}"
)
add_body("Newly blocked: kompass.com, europages.eu, ibisworld.com, statista.com, scribd.com, lusha.com, zoominfo.com, ukft.org, shopify.com, textileinfomedia.com, and 30+ more.")

add_body("**Fix B — Article/guide page filter:**")
add_body("Detects how-to articles, industry reports, and list pages by title keywords before they waste API credits:")
add_code_block(
    'ARTICLE_TITLE_SIGNALS = [\n'
    '  "top 10", "how to", "guide to", "industry analysis",\n'
    '  "2026 report", "directory of", "list of", "buyers & importers"...\n'
    ']\n'
    '// Blocks: "How to Source Fabric for a Clothing Line in 2026"\n'
    '// Keeps:  "Whaleys Bradford – Wholesale Fabric UK"'
)

add_body("**Fix C — Tier 3 Procurement Department Email Scraping:**")
add_body("When Snov.io and Hunter both find nothing, scrape the company website for a functional department email. Strict allowlist — no info@, no sales@:")
add_table(
    ["Accepted (strict allowlist)", "Rejected (everything else)"],
    [
        ["`procurement@`, `purchasing@`, `buying@`",                      "`info@`, `contact@`, `hello@`"],
        ["`sourcing@`, `import@`, `imports@`",                            "`sales@`, `support@`, `admin@`"],
        ["`einkauf@` (German: purchasing)",                                "`marketing@`, `webmaster@`"],
        ["`achat@`, `achats@` (French: purchasing)",                       "Anything not in the allowlist"],
        ["`inkoop@` (Dutch), `acquisti@` (Italian), `compras@` (Spanish)", ""],
    ]
)
add_body("Department emails reach the buying desk directly. In the results UI, department emails show a **dept.** badge to distinguish them from named person emails.")


# ══════════════════════════════════════════════════════════════════════════════
#  FINAL SECTION — TASK TRACKER
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()

# Big heading
tracker_heading = doc.add_heading("Task Tracker", level=0)
tracker_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in tracker_heading.runs:
    run.font.name  = "Calibri"
    run.font.size  = Pt(24)
    run.font.bold  = True
    run.font.color.rgb = DARK_BLUE
tracker_heading.paragraph_format.space_before = Pt(0)
tracker_heading.paragraph_format.space_after  = Pt(4)

subhead = doc.add_paragraph("Harsh Patel & Prateek — Agent 1 Development Progress")
r = subhead.runs[0] if subhead.runs else subhead.add_run("Harsh Patel & Prateek — Agent 1 Development Progress")
r.font.name  = "Calibri"
r.font.size  = Pt(13)
r.font.color.rgb = RGBColor(0x60, 0x7D, 0xA8)
subhead.paragraph_format.space_after = Pt(16)

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(0)
note.paragraph_format.space_after  = Pt(0)
nr = note.add_run("[ Insert tracker screenshot below ]")
nr.font.name  = "Calibri"
nr.font.size  = Pt(10)
nr.font.color.rgb = RGBColor(0xAB, 0xB5, 0xC4)
nr.italic = True
note.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Image placeholder box
for _ in range(22):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EFF6FF")
    pPr.append(shd)

spacer(10)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(DOCX_FILE)
print(f"✅  Saved: {DOCX_FILE}")
