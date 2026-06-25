#!/usr/bin/env python3
"""
Generate Word documents from Claude project export data.
Creates one .docx per project with: metadata, docs, and date-range conversations.
"""

import json
import os
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/harshpatel/Desktop/Elan-Exports_CRM/data-6d74feb6-cd86-4d3b-a609-ee0f3f215078-1781522434-c91936d3-batch-0000"
OUT_DIR = "/Users/harshpatel/Desktop/Elan-Exports_CRM"

# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc):
    """Add a horizontal line separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_colored_heading(doc, text, level, color_hex="1F497D"):
    """Add a heading with custom color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    return h

def add_info_box(doc, label, value, label_color="2E74B5", bg_color="EBF3FB"):
    """Add a label-value info row with background color."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    # remove borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    c0, c1 = table.columns[0].width, table.columns[1].width
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(13)

    cell_l = table.rows[0].cells[0]
    cell_r = table.rows[0].cells[1]
    set_cell_bg(cell_l, "D6E4F0")
    set_cell_bg(cell_r, bg_color)

    p_l = cell_l.paragraphs[0]
    run_l = p_l.add_run(label)
    run_l.bold = True
    run_l.font.size = Pt(10)
    run_l.font.color.rgb = RGBColor.from_string(label_color)

    p_r = cell_r.paragraphs[0]
    run_r = p_r.add_run(str(value))
    run_r.font.size = Pt(10)

    doc.add_paragraph()  # spacing

def add_message_block(doc, sender, text, timestamp):
    """Add a formatted conversation message block."""
    is_human = sender == "human"

    # Sender header
    sender_p = doc.add_paragraph()
    sender_p.paragraph_format.space_before = Pt(8)
    sender_p.paragraph_format.space_after = Pt(2)
    if is_human:
        sender_p.paragraph_format.left_indent = Inches(0)
        run = sender_p.add_run("  YOU  ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        # Add background via highlight workaround - use table instead
    else:
        sender_p.paragraph_format.left_indent = Inches(0)
        run = sender_p.add_run("  CLAUDE  ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Use a 1-row table for colored background message
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]

    if is_human:
        set_cell_bg(cell, "E8F4FD")
        # Remove border and add left accent
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'right', 'bottom', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:color'), '2E74B5')
        tcBorders.append(left)
        tcPr.append(tcBorders)
    else:
        set_cell_bg(cell, "F0F9F0")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'right', 'bottom', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:color'), '27AE60')
        tcBorders.append(left)
        tcPr.append(tcBorders)

    # Sender label inside cell
    p_header = cell.paragraphs[0]
    label_run = p_header.add_run(f"{'YOU' if is_human else 'CLAUDE'}   {timestamp}")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor.from_string("2E74B5" if is_human else "1E8449")

    # Message text
    lines = text.strip().split('\n')
    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_cover_page(doc, project):
    """Build the cover page for the project."""
    # Title
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(project['name'])
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor.from_string("1F497D")

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Claude AI Project Export")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor.from_string("666666")
    sub_run.italic = True

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    # Metadata table
    created = project.get('created_at', '')[:10]
    updated = project.get('updated_at', '')[:10]
    description = project.get('description', '') or 'No description'
    docs_count = len(project.get('docs', []))
    conv_count = project.get('_conv_count', 0)
    msg_count = project.get('_msg_count', 0)

    info_items = [
        ("Project Name", project['name']),
        ("Created", created),
        ("Last Updated", updated),
        ("Knowledge Base Files", str(docs_count)),
        ("Conversations Included", str(conv_count)),
        ("Total Messages", str(msg_count)),
        ("Exported By", "Elan (sales@elanexports.com)"),
        ("Export Date", "2026-06-18"),
    ]

    table = doc.add_table(rows=len(info_items), cols=2)
    table.style = 'Table Grid'

    for i, (label, val) in enumerate(info_items):
        row = table.rows[i]
        set_cell_bg(row.cells[0], "D6E4F0")
        set_cell_bg(row.cells[1], "EBF3FB")

        p_l = row.cells[0].paragraphs[0]
        r = p_l.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("1F497D")

        p_r = row.cells[1].paragraphs[0]
        p_r.add_run(val).font.size = Pt(10)

    doc.add_page_break()


def build_project_info(doc, project):
    """Build project description and system prompt section."""
    add_colored_heading(doc, "Project Overview", 1, "1F497D")
    add_horizontal_rule(doc)

    description = project.get('description', '') or 'No description provided.'
    desc_p = doc.add_paragraph(description)
    desc_p.paragraph_format.space_before = Pt(6)
    desc_p.paragraph_format.space_after = Pt(12)
    for run in desc_p.runs:
        run.font.size = Pt(11)

    prompt = project.get('prompt_template', '').strip()
    if prompt:
        add_colored_heading(doc, "System Prompt (Custom Instructions)", 2, "2E74B5")
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        set_cell_bg(cell, "FFF8DC")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'right', 'bottom', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        left_b = OxmlElement('w:left')
        left_b.set(qn('w:val'), 'single')
        left_b.set(qn('w:sz'), '12')
        left_b.set(qn('w:color'), 'E67E22')
        tcBorders.append(left_b)
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        r = p.add_run(prompt)
        r.font.size = Pt(10)
        r.font.name = "Courier New"
        doc.add_paragraph()


def build_docs_section(doc, project):
    """Build knowledge base documents section."""
    docs = project.get('docs', [])
    if not docs:
        add_colored_heading(doc, "Knowledge Base", 1, "1F497D")
        add_horizontal_rule(doc)
        p = doc.add_paragraph("No knowledge base files were uploaded to this project.")
        p.runs[0].font.color.rgb = RGBColor.from_string("888888")
        p.runs[0].italic = True
        return

    add_colored_heading(doc, f"Knowledge Base ({len(docs)} file{'s' if len(docs)!=1 else ''})", 1, "1F497D")
    add_horizontal_rule(doc)

    for i, doc_item in enumerate(docs):
        filename = doc_item.get('filename', f'Document {i+1}')
        content = doc_item.get('content', '')
        created = doc_item.get('created_at', '')[:10] if doc_item.get('created_at') else ''

        # File header
        file_heading = doc.add_paragraph()
        file_heading.paragraph_format.space_before = Pt(12)
        icon_run = file_heading.add_run(f"📄  {filename}")
        icon_run.bold = True
        icon_run.font.size = Pt(13)
        icon_run.font.color.rgb = RGBColor.from_string("2E74B5")

        if created:
            meta_p = doc.add_paragraph(f"Uploaded: {created}  |  Size: {len(content):,} characters")
            meta_p.runs[0].font.size = Pt(9)
            meta_p.runs[0].font.color.rgb = RGBColor.from_string("888888")

        # Content in bordered box
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        set_cell_bg(cell, "F8F9FA")

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'right', 'bottom']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), 'CCCCCC')
            tcBorders.append(border)
        tcPr.append(tcBorders)

        lines = content.strip().split('\n')
        first = True
        for line in lines:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.add_run(line).font.size = Pt(10)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        doc.add_paragraph()

    doc.add_page_break()


def build_conversations_section(doc, conversations, date_range_note):
    """Build the conversations section."""
    total_msgs = sum(len(c.get('chat_messages', [])) for c in conversations)

    add_colored_heading(doc, f"Conversations ({len(conversations)} conversations, {total_msgs} messages)", 1, "1F497D")

    # Date range note
    note_p = doc.add_paragraph()
    note_run = note_p.add_run(f"Note: {date_range_note}")
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor.from_string("888888")
    note_p.paragraph_format.space_after = Pt(8)

    add_horizontal_rule(doc)

    if not conversations:
        p = doc.add_paragraph("No conversations found for this project's date range.")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor.from_string("888888")
        return

    for conv_idx, conv in enumerate(conversations):
        name = conv.get('name', '') or conv.get('summary', '') or f"Conversation {conv_idx + 1}"
        created = conv.get('created_at', '')[:10]
        msgs = conv.get('chat_messages', [])

        # Conversation header
        conv_header = doc.add_paragraph()
        conv_header.paragraph_format.space_before = Pt(16)
        h_run = conv_header.add_run(f"{'─'*3}  {name}")
        h_run.bold = True
        h_run.font.size = Pt(12)
        h_run.font.color.rgb = RGBColor.from_string("34495E")

        meta_p = doc.add_paragraph(f"Date: {created}  |  Messages: {len(msgs)}")
        meta_p.runs[0].font.size = Pt(9)
        meta_p.runs[0].font.color.rgb = RGBColor.from_string("999999")

        # Filter messages with actual content
        real_msgs = [m for m in msgs if m.get('text', '').strip() or m.get('content')]

        if not real_msgs:
            p = doc.add_paragraph("(No message content available in export)")
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor.from_string("AAAAAA")
            p.runs[0].font.size = Pt(9)
            continue

        for msg in real_msgs:
            sender = msg.get('sender', 'unknown')
            text = msg.get('text', '').strip()
            ts = msg.get('created_at', '')[:16].replace('T', ' ')

            if not text:
                # Try content array
                content_blocks = msg.get('content', [])
                texts = []
                for block in content_blocks:
                    if isinstance(block, dict):
                        if block.get('type') == 'text':
                            texts.append(block.get('text', ''))
                        elif 'text' in block:
                            texts.append(block['text'])
                text = '\n'.join(texts).strip()

            if not text:
                continue

            # Truncate very long messages for readability
            if len(text) > 8000:
                text = text[:8000] + f"\n\n[... {len(text) - 8000:,} more characters truncated ...]"

            add_message_block(doc, sender, text, ts)

        add_horizontal_rule(doc)


# ── main ───────────────────────────────────────────────────────────────────────

def load_all_data():
    """Load all projects and conversations."""
    project_files = sorted(os.listdir(f"{BASE}/projects"))
    projects = []
    for pf in project_files:
        with open(f"{BASE}/projects/{pf}") as f:
            projects.append(json.load(f))

    with open(f"{BASE}/conversations.json") as f:
        conversations = json.load(f)

    return projects, conversations


def filter_conversations_by_date(conversations, start_date, end_date):
    """Filter conversations by creation date range (strings like '2026-02-25')."""
    result = []
    for conv in conversations:
        created = conv.get('created_at', '')[:10]
        if start_date <= created <= end_date:
            result.append(conv)
    return sorted(result, key=lambda c: c.get('created_at', ''))


def generate_doc(project, conversations, date_range_note, output_path):
    """Generate a Word document for a project."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Stats for cover
    project['_conv_count'] = len(conversations)
    project['_msg_count'] = sum(len(c.get('chat_messages', [])) for c in conversations)

    build_cover_page(doc, project)
    build_project_info(doc, project)
    build_docs_section(doc, project)
    build_conversations_section(doc, conversations, date_range_note)

    doc.save(output_path)
    print(f"✓ Saved: {output_path}")


def main():
    projects, all_conversations = load_all_data()

    # First 4 projects with their date ranges
    # Date ranges: from project creation until the next project's creation date - 1 day
    # Project dates: Feb25, Feb27, Apr27, Apr29, Apr29...
    config = [
        {
            "project": projects[0],  # How to use Claude (Feb 25)
            "start": "2026-02-25",
            "end": "2026-02-26",
            "note": "Conversations from Feb 25–26, 2026 (project creation period). The Claude export does not explicitly link conversations to projects, so these are matched by date."
        },
        {
            "project": projects[1],  # Tech Team Research (Feb 27)
            "start": "2026-02-27",
            "end": "2026-04-26",
            "note": "Conversations from Feb 27 – Apr 26, 2026 (Tech Team Research project period). Matched by date range between project creation dates."
        },
        {
            "project": projects[2],  # EEC Social Media Marketing Engine (Apr 27)
            "start": "2026-04-27",
            "end": "2026-04-28",
            "note": "Conversations from Apr 27–28, 2026 (EEC Social Media Marketing Engine creation period). Matched by date."
        },
        {
            "project": projects[3],  # Productivity (Apr 29)
            "start": "2026-04-29",
            "end": "2026-05-04",
            "note": "Conversations from Apr 29 – May 4, 2026 (Productivity project period). Matched by date."
        },
    ]

    for cfg in config:
        project = cfg["project"]
        convs = filter_conversations_by_date(all_conversations, cfg["start"], cfg["end"])

        # Sanitize project name for filename
        safe_name = project['name'].replace('/', '-').replace('\\', '-').replace(':', '-').replace('–', '-').strip()
        output_path = f"{OUT_DIR}/{safe_name}.docx"

        print(f"\nGenerating: {safe_name}")
        print(f"  Date range: {cfg['start']} → {cfg['end']}")
        print(f"  Conversations: {len(convs)}")

        generate_doc(project, convs, cfg["note"], output_path)


if __name__ == "__main__":
    main()
